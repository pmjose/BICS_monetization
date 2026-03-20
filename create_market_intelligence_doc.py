from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(30, 58, 95)

title = doc.add_heading('Market Intelligence Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Understanding the Telco Data Monetization Landscape')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph('BICS | Telco Data Monetization Platform').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('February 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

doc.add_heading('Table of Contents', 1)
toc_items = [
    ('1. Executive Summary', '3'),
    ('2. Belgium & EU Context: Digital Decade & Data Strategy', '4'),
    ('3. Belgian Smart City & IoT Opportunity', '6'),
    ('4. Data Product Trends', '8'),
    ('5. Market Demand', '11'),
    ('6. Pricing Models', '13'),
    ('7. Data Distribution Options', '15'),
    ('8. References', '17'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.add_run('\t' * 8 + page)

doc.add_page_break()

doc.add_heading('1. Executive Summary', 1)

doc.add_paragraph(
    'This Market Intelligence report addresses key strategic questions about telco data monetization, '
    'providing research, market data, and actionable insights for BICS.'
)

doc.add_heading('Key Questions Addressed', 2)

questions = [
    ('What is the telco data product trend?', 
     'Three major shifts: from raw data to privacy-safe insights, from one-off deals to marketplaces, and AI-native use on governed products.'),
    ('What is the market demand for telco data products?', 
     'Multi-billion dollar market (~$5.3B in 2024, growing to ~$14.9B by 2029 at ~24% CAGR) with strong demand from public sector, retail, transport, financial services, and advertising.'),
    ('Is there a reference pricing model?', 
     'A pricing toolbox: subscription/recurring licenses, usage-based/API-metered, flat licenses, freemium tiers, and outcome/revenue-share models.'),
    ('How to expose data to consumers without a Snowflake account?', 
     'Multiple options: Reader accounts, "Powered-by BICS" portals, REST APIs, Clean Rooms, and file exports - all with Snowflake as the governed backend.'),
]

for q, a in questions:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

doc.add_heading('2. Belgium & EU Context: Digital Decade & Data Strategy', 1)

doc.add_heading('EU Digital Decade & Data Act', 2)

doc.add_paragraph(
    'The European Commission\'s Digital Decade programme and the EU Data Act (2024) drive Europe\'s data agenda. '
    'Belgium, as host of EU institutions and home to BICS (Proximus Global), is uniquely positioned:'
)

phases = [
    ('2024 (EU Data Act)', 'New rules for fair access to and use of data across all economic sectors'),
    ('2025 (European Data Spaces)', 'Sector-specific data spaces for health, transport, energy, agriculture'),
    ('2030 (Digital Decade)', '75% of EU enterprises using cloud/AI/big data; 80% population with digital ID'),
]

for phase, desc in phases:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(phase + ' - ').bold = True
    p.add_run(desc)

doc.add_heading('Key Statistics', 2)

stats_table = doc.add_table(rows=4, cols=2)
stats_table.style = 'Table Grid'
stats_data = [
    ('Metric', 'Value'),
    ('EU Digital Economy Target', 'Top Global Digital Leader by 2030'),
    ('Belgian Telecom Market', '11.5M mobile subscribers'),
    ('BICS Global IoT Reach', '200+ countries, 700+ operators'),
]
for i, (metric, value) in enumerate(stats_data):
    row = stats_table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = value
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_heading('GDPR & Privacy Framework', 2)
doc.add_paragraph(
    'Belgium\'s Data Protection Authority (APD/GBA) enforces GDPR, making privacy-safe data products essential. '
    'BICS\'s approach to aggregated, anonymized insights is fully aligned with GDPR requirements, '
    'enabling compliant data monetization across EU markets.'
)

doc.add_page_break()

doc.add_heading('3. Belgian Smart City & IoT Opportunity', 1)

doc.add_heading('Belgian Smart City Initiatives', 2)

giga_table = doc.add_table(rows=5, cols=3)
giga_table.style = 'Table Grid'
giga_data = [
    ('Initiative', 'Description', 'Data Opportunity'),
    ('Smart Flanders', 'Digital transformation of Flemish cities', 'Mobility, parking, air quality analytics'),
    ('Brussels Smart City', 'Capital region digital infrastructure', 'Urban flow, transit, tourism analytics'),
    ('Port of Antwerp-Bruges', 'Europe\'s 2nd largest port, IoT-enabled', 'Container tracking, logistics optimization'),
    ('imec Research Hub', 'World-leading nanotech & IoT research', 'Edge AI, sensor analytics, 5G testbeds'),
]
for i, row_data in enumerate(giga_data):
    row = giga_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            row.cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_heading('BICS IoT Industry Verticals', 2)

iot_verticals = [
    ('Agriculture', 'Precision farming, soil monitoring, livestock tracking across Flanders & Wallonia'),
    ('Healthcare', 'Patient monitoring, asset tracking, cold chain at UZ hospitals across Belgium'),
    ('Industrial Manufacturing', 'Predictive maintenance, energy monitoring at Port of Antwerp & industrial zones'),
    ('Transport & Logistics', 'Fleet tracking, container monitoring at ports, airports & logistics hubs'),
    ('OEM', 'Connected devices, embedded eSIM/iSIM, edge compute for smart city deployments'),
]

for name, desc in iot_verticals:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ' - ').bold = True
    p.add_run(desc)

doc.add_heading('BICS Connectivity Services', 2)

connectivity = [
    ('Global IoT Connectivity', 'NB-IoT, LTE-M, 4G, 5G across 200+ countries via 700+ operator partners'),
    ('eSIM Hub', 'Remote SIM provisioning for IoT at scale'),
    ('IoT eSIM & iSIM', 'Embedded connectivity for OEM and industrial devices'),
    ('Cloud Connectivity', 'Direct cloud interconnect for IoT data pipelines'),
    ('IoT Security Solutions', 'End-to-end security for connected devices'),
    ('Advanced Analytics', 'Real-time telemetry analytics and alerting platform'),
]

for name, desc in connectivity:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ' - ').bold = True
    p.add_run(desc)

doc.add_heading('Belgian Tourism & Events', 2)

tourism = [
    ('Brussels EU Quarter', 'Visitor flow analytics for EU institution district'),
    ('Bruges Historic Center', 'UNESCO site crowd management and seasonal patterns'),
    ('Tomorrowland / Festivals', 'Event crowd management, transport optimization'),
    ('Ardennes Tourism', 'Nature tourism flow analytics and seasonal planning'),
    ('Belgian Coast', 'Coastal tourism density and transport patterns'),
]

for name, desc in tourism:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(name + ' - ').bold = True
    p.add_run(desc)

doc.add_page_break()

doc.add_heading('Data & AI Governance Framework', 2)

governance = [
    ('GDPR / APD-GBA', ['Data Protection Impact Assessments', 'Consent management', 'Right to erasure', 'Privacy-by-design']),
    ('EU Data Act (2024)', ['Fair data access rules', 'B2B/B2G data sharing', 'Cloud switching rights', 'Interoperability standards']),
    ('EU AI Act (2024)', ['Risk-based AI classification', 'High-risk AI requirements', 'Transparency obligations', 'Conformity assessments']),
]

for name, items in governance:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('4. Data Product Trends', 1)

doc.add_heading('What is a Data Product?', 2)

doc.add_paragraph(
    'A Data Product is a curated, governed, and reusable collection of data assets designed to deliver '
    'specific value to a defined set of consumers. Unlike raw data exports, a data product is:'
)

dp_chars = [
    ('Self-describing', 'metadata, documentation, quality metrics'),
    ('Access-controlled', 'built-in governance and privacy'),
    ('Consumption-ready', 'optimized for specific use cases'),
]

for char, desc in dp_chars:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(char + ' - ').bold = True
    p.add_run(desc)

doc.add_heading('What This Means in the Telco World', 2)

telco_meaning = [
    ('Not raw CDRs or network logs', 'Aggregated, anonymized insights (mobility, footfall, dwell times)'),
    ('Not one-off data dumps', 'Continuously refreshed feeds with SLAs & quality monitoring'),
    ('Not "take it or leave it"', 'Tiered products by geography, granularity, history depth'),
    ('Not uncontrolled sharing', 'Privacy-safe delivery via clean rooms & differential privacy'),
    ('Not just data', 'Data + AI models + APIs as a service'),
]

for not_this, instead in telco_meaning:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(not_this + ' -> ').bold = True
    p.add_run(instead)

doc.add_heading('Three Big Shifts in Telco Data Products', 2)

shifts = [
    ('1. From Raw Data to Privacy-Safe Insight Products',
     'Operators are moving away from selling raw CDR/location feeds toward curated, anonymized insight products: '
     'mobility & footfall, network quality, IoT/5G telemetry, fraud/risk signals, customer intent, and more.',
     'Example: "Digital nervous system" products for cities - real-time density, surge prediction, and recommended '
     'actions for transport, energy, and public safety packaged as data products, not raw network logs.'),
    ('2. From One-Off Deals to Marketplaces & Platforms',
     'Telcos increasingly publish their data products through data marketplaces and clean rooms, so partners '
     '(cities, retailers, banks, advertisers) can subscribe to governed feeds and run their own analytics without '
     'ever seeing raw subscriber identifiers.',
     'Key principle: "This is not about selling raw telco data, it\'s about a trusted collaboration platform for insights."'),
    ('3. AI-Native / Agentic Use on Governed Products',
     'The next wave is AI/agent-assist on top of these data products: call-center intent detection, agent assist, '
     'network ops copilots - all using the same governed products with masking and row-level policies.',
     'The pattern: Build reusable, policy-attached telco data products once; reuse them across dashboards, AI, and partner use cases.'),
]

for title, desc, example in shifts:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(desc)
    p2 = doc.add_paragraph()
    p2.add_run(example).italic = True
    doc.add_paragraph()

doc.add_heading('Data Product Categories', 2)

categories_table = doc.add_table(rows=7, cols=3)
categories_table.style = 'Table Grid'
cat_data = [
    ('Category', 'Maturity', 'Privacy Model'),
    ('Mobility & Footfall', '95%', 'Aggregated'),
    ('Audience Segments', '85%', 'Cohort-based'),
    ('Network Quality', '80%', 'Aggregated'),
    ('Fraud & Risk Signals', '75%', 'Scored'),
    ('IoT / 5G Telemetry', '70%', 'Device-level'),
    ('Customer Intent', '65%', 'Inferred'),
]
for i, row_data in enumerate(cat_data):
    row = categories_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            row.cells[j].paragraphs[0].runs[0].bold = True

doc.add_page_break()

doc.add_heading('TM Forum Industry Guidance', 2)

doc.add_paragraph(
    'TM Forum, the global industry association for digital business, provides comprehensive guidance on telco '
    'data monetization through its Open Digital Framework and research publications.'
)

doc.add_heading('IG1138: External Data Monetization Guide', 3)
ig1138 = [
    'Regulations & Compliance - Privacy requirements by jurisdiction',
    'Privacy Techniques - Anonymization, aggregation, differential privacy',
    '10 Monetization Use Cases across Health, Transportation, Finance/Fraud, Advertising, Site Planning, and Tourism',
    'Revenue Conversion - Moving from data assets to commercial products',
]
for item in ig1138:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('TM Forum Use Case Framework', 3)

usecase_table = doc.add_table(rows=7, cols=3)
usecase_table.style = 'Table Grid'
usecase_data = [
    ('Vertical', 'Data Products', 'Privacy Model'),
    ('Health', 'Population health patterns, mobility for outbreak tracking', 'Aggregated'),
    ('Transportation', 'Traffic flow, commuter patterns, route optimization', 'Aggregated'),
    ('Finance/Fraud', 'Fraud detection signals, credit risk scoring, geo-behavior', 'Scored/Anonymized'),
    ('Advertising', 'Audience segments, campaign measurement, attribution', 'Cohort-based'),
    ('Site Planning', 'Footfall analytics, catchment analysis, location scoring', 'Aggregated'),
    ('Tourism', 'Visitor flows, origin-destination, seasonal patterns', 'Aggregated'),
]
for i, row_data in enumerate(usecase_data):
    row = usecase_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            row.cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph('Source: TM Forum IG1138 - Introductory Guide to External Data Monetization').italic = True

doc.add_heading('GSMA Data Monetization Insights', 2)

doc.add_paragraph(
    'GSMA Intelligence, the research arm of the global mobile industry association representing operators worldwide, '
    'provides critical market intelligence on telco data monetization opportunities.'
)

gsma_points = [
    ('B2B Market Opportunity', '$400B addressable market for B2B services beyond connectivity'),
    ('Enterprise Digital Transformation', '10% of operator revenue (2024-2030)'),
    ('IoT Connections', '26 billion globally by 2025'),
    ('AI Monetization', '74% of operators testing generative AI'),
    ('Open Gateway', '74% of operators committed to API exposure'),
]

for title, desc in gsma_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

doc.add_heading('5. Market Demand', 1)

doc.add_heading('Market Size & Growth', 2)

market_table = doc.add_table(rows=4, cols=2)
market_table.style = 'Table Grid'
market_data = [
    ('Metric', 'Value'),
    ('Global Market 2024', '~$5.3B'),
    ('Projected by 2029', '$14.9B'),
    ('CAGR Growth Rate', '~24%'),
]
for i, (metric, value) in enumerate(market_data):
    row = market_table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = value
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Message: ').bold = True
p.add_run('"This is now a multi-billion-dollar, high-growth segment, not an experiment."')

doc.add_heading('Market Projections (2024-2029)', 2)

proj_table = doc.add_table(rows=7, cols=2)
proj_table.style = 'Table Grid'
proj_data = [
    ('Year', 'Market Size ($B)'),
    ('2024', '5.3'),
    ('2025', '6.6'),
    ('2026', '8.2'),
    ('2027', '10.2'),
    ('2028', '12.5'),
    ('2029', '14.9'),
]
for i, (year, size) in enumerate(proj_data):
    row = proj_table.rows[i]
    row.cells[0].text = year
    row.cells[1].text = size
    if i == 0:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph('Sources: Precedence Research (2024), Market Research Future').italic = True

doc.add_heading('Who Buys Telco Data Products?', 2)

buyers = [
    ('Government / Public Sector', 'Urban planning, transit optimization, emergency response, tourism boards'),
    ('Retail & CPG', 'Site selection, trade area analysis, campaign measurement, competitive intelligence'),
    ('Financial Services', 'Fraud signals, credit scoring, footfall proxies for revenue, insurance risk'),
    ('Transportation & Logistics', 'Route planning, congestion analytics, fleet optimization'),
    ('Real Estate & Property', 'Location scoring, development feasibility, tenant analytics'),
    ('Advertising & Media', 'Audience segments, attribution, OOH measurement'),
]

for buyer, uses in buyers:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(buyer + ': ').bold = True
    p.add_run(uses)

doc.add_page_break()

doc.add_heading('6. Pricing Models', 1)

doc.add_paragraph(
    'There is no single "industry standard" price list for telco data products. Instead, there is a toolbox '
    'of pricing models that telcos mix and match based on product type, buyer sophistication, and commercial objectives.'
)

doc.add_heading('Pricing Model Options', 2)

pricing_models = [
    ('Subscription / Recurring License',
     'Monthly or annual fee for ongoing access to a data feed or dashboard',
     'Best For: Continuous use cases (e.g., weekly footfall reports)',
     'Example: $5K-$50K/month for city-level mobility feed'),
    ('Usage-Based / API Metered',
     'Pay-per-query, per-record, or per-API-call',
     'Best For: Variable demand, developer ecosystems',
     'Example: $0.001-$0.10 per API call or per 1K records'),
    ('Flat License / One-Time',
     'Single payment for a defined dataset or historical archive',
     'Best For: One-off studies, academic research',
     'Example: $10K-$500K for historical mobility archive'),
    ('Freemium / Tiered',
     'Free tier with limited scope; paid tiers unlock geography, granularity, or history',
     'Best For: Land-and-expand, developer adoption',
     'Example: Free city-level aggregates; paid for district-level or hourly'),
    ('Outcome / Revenue Share',
     'Pricing tied to buyer\'s business outcome (e.g., % of ad spend, % of revenue uplift)',
     'Best For: High-trust partnerships, joint ventures',
     'Example: 5-15% of incremental revenue attributed to data'),
]

for name, desc, best_for, example in pricing_models:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    doc.add_paragraph(desc)
    doc.add_paragraph(best_for, style='List Bullet')
    doc.add_paragraph(example, style='List Bullet')
    doc.add_paragraph()

doc.add_heading('Pricing Recommendations', 2)

recommendations = [
    'Start with subscription for predictable revenue and customer stickiness',
    'Offer usage-based for API products to lower entry barrier',
    'Use tiered pricing to capture different willingness-to-pay segments',
    'Reserve outcome-based for strategic partnerships with shared upside',
    'Always include governance and audit trail as value-add, not cost',
]

for rec in recommendations:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_page_break()

doc.add_heading('7. Data Distribution Options', 1)

doc.add_paragraph(
    'How to expose data to consumers without requiring them to have a Snowflake account? Multiple options exist, '
    'all with Snowflake as the governed backend.'
)

doc.add_heading('Distribution Channels', 2)

distribution = [
    ('Snowflake Reader Accounts',
     'Snowflake-hosted accounts for partners who don\'t have their own Snowflake. They query your shares using Snowflake UI or connectors.',
     'Pros: Full SQL power, governed, audit trail',
     'Cons: Requires Snowflake literacy; you pay compute'),
    ('"Powered by BICS" Portal / Streamlit',
     'Custom web app (e.g., Streamlit) that queries Snowflake and presents insights via UI. No Snowflake login needed.',
     'Pros: Branded experience, no Snowflake knowledge required',
     'Cons: Development effort; limited to pre-built views'),
    ('REST APIs',
     'Expose Snowflake data via REST endpoints (e.g., using External Functions, Snowpark Container Services, or a gateway).',
     'Pros: Universal integration, developer-friendly',
     'Cons: Requires API management; usage metering'),
    ('Clean Rooms (Snowflake Data Clean Rooms)',
     'Privacy-safe collaboration where partners bring their data into a governed environment for joint analysis without raw data exposure.',
     'Pros: Privacy-compliant, high-value partnerships',
     'Cons: Requires partner engagement; more complex setup'),
    ('File Exports (S3, SFTP, etc.)',
     'Scheduled exports of aggregated data to cloud storage or SFTP for partners with legacy systems.',
     'Pros: Works with any system',
     'Cons: No real-time; less governed; data leaves your control'),
]

for name, desc, pros, cons in distribution:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    doc.add_paragraph(desc)
    doc.add_paragraph(pros, style='List Bullet')
    doc.add_paragraph(cons, style='List Bullet')
    doc.add_paragraph()

doc.add_heading('Distribution Recommendations', 2)

dist_recs = [
    'Default to Reader Accounts or Streamlit for most use cases',
    'Use REST APIs for developer ecosystems and programmatic access',
    'Reserve Clean Rooms for high-value, privacy-sensitive partnerships',
    'Avoid raw file exports unless absolutely necessary (regulatory, legacy systems)',
    'Always maintain Snowflake as the single source of truth for governance',
]

for rec in dist_recs:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_page_break()

doc.add_heading('8. References', 1)

doc.add_heading('TM Forum', 2)
tmf_refs = [
    'GB1086: Data Product Lifecycle Management (DPLM) - Establishing an Operational Model for Telco Data v1.0.0',
    'IG1138: Introductory Guide to External Data Monetization R15.5.1',
    'ODA Monetization Engine Catalyst (2024)',
    'TM Forum Inform: "How telcos can monetize data through digital marketplaces"',
]
for ref in tmf_refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('GSMA', 2)
gsma_refs = [
    'The Mobile Economy 2025',
    'Operator Strategies for Monetization and Customer Experience',
    'GSMA Open Gateway Initiative',
    'GSMA Intelligence Research',
]
for ref in gsma_refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('Belgium & EU', 2)
eu_refs = [
    'European Commission - Digital Decade (digital-strategy.ec.europa.eu)',
    'EU Data Act (2024)',
    'EU AI Act (2024)',
    'Belgian Data Protection Authority - APD/GBA (dataprotectionauthority.be)',
    'BICS - IoT Connectivity (bics.com)',
    'Proximus Group (proximus.com)',
    'Smart Flanders (smart.flanders.be)',
    'imec Research (imec-int.com)',
]
for ref in eu_refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_heading('Market Research', 2)
market_refs = [
    'Precedence Research - Telecom Analytics Market Report (2024)',
    'Market Research Future - Telecom Data Monetization Forecast',
]
for ref in market_refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('--- End of Report ---').italic = True

output_path = '/Users/pjose/Documents/GitHub/BICS_monetization/BICS_Market_Intelligence_Report.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
